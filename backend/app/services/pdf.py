import io
from typing import List, Dict
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

def notes_to_pdf(notes: List[Dict[str, str]]) -> io.BytesIO:
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    def new_page_header():
        p.setFont("Helvetica-Bold", 16)
        p.drawString(40, height - 40, "CurioMindAI Notes")
        p.setFont("Helvetica", 12)

    new_page_header()
    y = height - 70

    for i, note in enumerate(notes):
        # Wrap long text manually (very simple wrap)
        q = f"Q{i+1}: {note['question']}"
        e = note['explanation']
        lines = [q] + [e]
        for line in lines:
            if y < 60:
                p.showPage()
                new_page_header()
                y = height - 70
            p.drawString(40, y, line)
            y -= 18
        y -= 12

    p.save()
    buffer.seek(0)
    return buffer

def notes_to_docx(notes: List[Dict[str, str]]) -> io.BytesIO:
    doc = Document()
    doc.add_heading('CurioMindAI Notes', level=1)
    for i, note in enumerate(notes, start=1):
        doc.add_heading(f"Q{i}: {note['question']}", level=3)
        p = doc.add_paragraph(note['explanation'])
        for run in p.runs:
            run.font.size = Pt(11)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def make_suggestions(seed: str) -> List[str]:
    seed = (seed or '').strip()
    if not seed:
        return [
            'What are black holes?',
            'How does photosynthesis work?',
            'What is the water cycle?',
            'How does gravity affect us?'
        ]

    # Clean up the seed into a concise topic fragment
    phrase = ' '.join(seed.split()[:6]).strip().rstrip('?.!')
    base = phrase[0:1].upper() + phrase[1:]
    candidates = [
        f'How is {base} used in real life?',
        f'Why is {base} important?',
        f'Simple example of {base}',
        f'How does {base} compare to similar ideas?',
        f'Where do we see {base} in everyday life?'
    ]
    # Return between 3 and 5 items
    return candidates[:5]
